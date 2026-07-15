# -*- coding: utf-8 -*-
"""Generate the moso-bamboo DPP4 inhibitory-peptide project-summary docx."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Repo root = two levels above this script (scripts/reports/); clone-able to regenerate independently
ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

doc = Document()

# Default font size
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

def h1(t):
    p = doc.add_heading(t, level=1)
    return p

def h2(t):
    p = doc.add_heading(t, level=2)
    return p

def para(t, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = bold
    return p

def bullet(t):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(t)
    return p

def numbered(t):
    p = doc.add_paragraph(style='List Number')
    p.add_run(t)
    return p

# ===== Title =====
title = doc.add_heading('Moso Bamboo DPP4 Inhibitory Peptides Project: Completed-Work Overview', level=0)
sub = doc.add_paragraph('Generated: 2026-07-14 update  |  Repo: github.com/wu-yijing/moso-bamboo-DPP4-peptides')
sub.runs[0].italic = True

# ===== I. Project goal and reasons for redirection =====
h1('I. Project goal and reasons for redirection')
para('The original yam (Dioscorea polystachya) project collapsed because UniProt had only 20 curated proteins, and re-running virtual digestion locally yielded only 237 peptides (breaking the original 5,230 funnel); moreover 0/7 of the experimentally-validated peptides (FWPQY etc.) could be cleaved from those 20 proteins - the yam could not support a DPP4 claim. We therefore turned to the reference template paper "Discovery of garlic-derived peptides as natural DPP4 inhibitors" (Cheng et al., Bioorganic Chemistry 175, 2026, 109801) and reselected moso bamboo (Moso bamboo) as the study object.')

# ===== II. Step-by-step details =====
h1('II. Step-by-step details (results / software / method)')

# Step 1
h2('Step 1 - Species identification and protein-pool verification')
bullet('Result: moso bamboo = Phyllostachys edulis, UniProt taxonomy 38705; curated (Swiss-Prot) = 0 entries, all entries (incl. TrEMBL) = 253 entries, no reference proteome; average 409 aa, total residues 103,409. vs the template garlic (A. sativum) pool of only 113 entries -> moso pool ~2.2x garlic, passes.')
bullet('Software: curl + UniProt REST API (taxonomy/search, uniprotkb/search, proteomes/search).')
bullet('Method: first reverse-lookup the species ID from taxonomy (do not trust colloquial names), then count reviewed vs all entries separately; benchmark against the template paper\'s actual usage.')

# Step 2
h2('Step 2 - Protein sequence download')
bullet('Result: moso_253.fasta (131 KB, 253 real sequences, 0 placeholders).')
bullet('Software: curl + UniProtKB /stream (FASTA format).')
bullet('Method: query=taxonomy_id:38705, full-library entry pool (not reviewed-only, consistent with the template paper\'s actual practice).')

# Step 3
h2('Step 3 - Simulated gastrointestinal digestion (virtual enzymatic hydrolysis)')
bullet('Result: strict rule (chymotrypsin cuts F/Y/W): 7,988 unique peptides, of which 4,950 are 2-6 aa short peptides; relaxed rule (add cut at L): 7,472 unique peptides, 4,761 2-6 aa. Scale = 3.4x the garlic template (1,442 short peptides).')
bullet('Software: Python (venv: numpy/scipy/biopython) + self-written script rerun_digestion_moso253_strict.py.')
bullet('Method: replicate ExPASy PeptideCutter - pepsin (pH1.3) + trypsin + specific chymotrypsin (cuts C-term of F/Y/W, no cut before Pro), merge and dedupe, keep 2-6 aa.')

# Step 4
h2('Step 4 - PeptideRanker-style scoring + allergen/toxicity filtering')
bullet('Result: >0.5 candidates 4,333 -> de-allergen (AllerTOP-style) / de-toxic (ToxinPred-style) -> 4,289.')
bullet('Software: Python (self-written proxy heuristic).')
bullet('Method: early transparent proxy scoring (based on literature physicochemical features: N-terminal hydrophobicity, Pro content, molecular weight). Because the PeptideRanker official server was long unavailable, and this proxy scored rich-Pro short peptides constantly at 1.000 with zero discrimination, it was replaced by the "iDPPIV-SCM local offline reproduction" (see Section VII); the allergen/toxicity layer is replaced by ToxinPred 3.0 + AlgPred 2.0 official outputs.')

# Step 5
h2('Step 5 - DPP4 structural-preference narrowing + docking queue')
bullet('Result: 3-5 aa, N-terminal hydrophobic, position-2 Pro/Ala preference -> candidate pool 2,019; take top-60 into the docking queue.')
bullet('Software: Python self-written moso_pipeline_filter2.py.')
bullet('Method: second-pass filtering per known DPP4 inhibitory-peptide structural rules (short peptide, N-terminal hydrophobic, Pro-enriched).')

# Step 6
h2('Step 6 - Ligand 3D structure preparation (PDBQT)')
bullet('Result: 60/60 ligand PDBQTs all generated successfully (3-7 KB each, zero NaN).')
bullet('Software: OpenBabel (pybel / obabel CLI) + RDKit (rescue 8 Arg/His peptides).')
bullet('Method: main SMILES -> make3D -> PDBQT; 8 Arg/His peptides produced NaN due to missing force-field data -> use RDKit MolFromSequence + AddHs + ETKDGv3 embed + MMFFOptimizeMolecule optimize -> convert to PDBQT.')

# Step 7
h2('Step 7 - Receptor and pocket preparation')
bullet('Result: 1WCY_receptor.pdbqt (12,248 atoms, correct format); pocket box moso_box.txt center (62.8, 47.7, 4.8), size 30^3.')
bullet('Software: RCSB PDB download + OpenBabel (obabel CLI) + awk fixed-column parse.')
bullet('Method: download DPP4 crystal 1WCY (sitagliptin-bound state) -> convert to PDBQT -> keep only ATOM/TER/END (strip HEADER etc. headers) -> measured ligand A1201 coordinates set the grid center (the template paper\'s given (54,62,37) is a same-pocket approximation; measured center is more robust).')

# Step 8
h2('Step 8 - Molecular docking (AutoDock Vina)')
para('Result: 60/60 completed. Top 10 (binding free energy dG, kcal/mol):', bold=True)

# Table
tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Light Grid Accent 1'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
for i, t in enumerate(['Rank', 'Peptide', 'dG (kcal/mol)', 'Class']):
    hdr[i].text = t
    hdr[i].paragraphs[0].runs[0].bold = True
rows = [
    ('1', 'LPPQ (Leu-Pro-Pro-Gln)', '-7.472', 'medium-strong'),
    ('2', 'APSPE (Ala-Pro-Ser-Pro-Glu)', '-7.150', 'medium-strong'),
    ('3', 'LAPSP (Leu-Ala-Pro-Ser-Pro)', '-7.087', 'medium-strong'),
    ('4', 'LPGP (Leu-Pro-Gly-Pro)', '-7.075', 'medium-strong'),
    ('5', 'LPINP (Leu-Pro-Ile-Asn-Pro)', '-6.988', 'medium-strong'),
    ('6', 'LPSP (Leu-Pro-Ser-Pro)', '-6.867', 'medium-strong'),
    ('7', 'LPCPR (Leu-Pro-Cys-Pro-Arg)', '-6.835', 'medium-strong'),
    ('8', 'LPGDP (Leu-Pro-Gly-Asp-Pro)', '-6.793', 'medium-strong'),
    ('9', 'LPDDP (Leu-Pro-Asp-Asp-Pro)', '-6.693', 'medium-strong'),
    ('10', 'APSQP (Ala-Pro-Ser-Gln-Pro)', '-6.515', 'medium-strong'),
]
for r in rows:
    c = tbl.add_row().cells
    for i, v in enumerate(r):
        c[i].text = v

para('')
bullet('Distribution: medium-strong (-6.5 ~ -8) 10 / medium (-5 ~ -6.5) 49 / weak (>-5) 1 (CPPSK -4.856); no strong binding (<-8). Champion LPPQ (Leu-Pro-Pro-Gln, double Pro) fully matches the literature signature of DPP4 inhibitory peptides, dG magnitude comparable to template garlic peptides WPHY/WPQY.')
bullet('Software: AutoDock Vina 1.2.5 (vina.exe, local Windows).')
bullet('Method: --exhaustiveness 4 --cpu 2, single docking box 30^3, 9 poses/peptide output, take best mode dG.')

# ===== III. Deliverables list =====
h1('III. Current deliverables list (data/ and docking/ in repo)')
for f in [
    'moso_253.fasta (253 proteins)',
    'moso_253_peptides_strict.txt / _peptides.txt (unique-peptide lists)',
    'moso_candidates_pr_filtered.txt (2,019 candidates)',
    'moso_dock_queue.txt (60-peptide docking queue)',
    '1WCY_receptor.pdbqt, moso_box.txt, vina.exe',
    'moso_ligands/ (60 ligand PDBQTs + 60 docked output poses)',
    'moso_dock_results.tsv, moso_dock_ranking.txt (full ranking)',
]:
    bullet(f)

# ===== IV. Pending =====
h1('IV. Remaining follow-up steps (full template pipeline)')
t2 = doc.add_table(rows=1, cols=3)
t2.style = 'Light Grid Accent 1'
h = t2.rows[0].cells
for i, t in enumerate(['Stage', 'Status', 'Note']):
    h[i].text = t
    h[i].paragraphs[0].runs[0].bold = True
for r in [
    ('Official PeptideRanker/AllerTOP/ToxinPred validation', 'Not done', 'currently proxy heuristic, must replace'),
    ('MD + MM/PBSA (GROMACS)', 'Not done', 'template uses 50-150 ns to validate Top peptides'),
    ('Network pharmacology', 'Not done', 'SwissTargetPrediction->STRING->DAVID'),
    ('In-vitro activity (Gly-Pro-pNA + Caco-2)', 'Not done', 'true activity endpoint, needs experiment'),
]:
    c = t2.add_row().cells
    for i, v in enumerate(r):
        c[i].text = v

# ===== V. Limitations =====
h1('V. Limitations that must be disclosed truthfully')
bullet('Activity pre-screening scoring is replaced by the iDPPIV-SCM local offline reproduction (literature-validated, offline-reproducible, interpretable) instead of the original proxy heuristic; the allergen/toxicity layer uses ToxinPred 3.0 + AlgPred 2.0 official outputs. The iDPPIV-SCM training benchmark has length confounding (negative samples mostly long peptides/proteins), so it is used only for candidate ranking/soft filtering, not deterministic activity judgment.')
bullet('Docking is a static dG estimate that needs MD/MM-PBSA + in-vitro experiments for confirmation; cannot be taken directly as an activity conclusion.')
bullet('Moso bamboo has 0 manually-reviewed entries (all TrEMBL predictions); Methods must truthfully disclose the source and predictive nature (the garlic template is also ~95% TrEMBL, acceptable).')
bullet('WPHY/WPQY/VAPGW are garlic peptides and must not be used for moso true-value validation.')

# ===== VI. Software stack =====
h1('VI. Project software stack summary')
for s in [
    'UniProt REST API (curl): taxonomy/search, uniprotkb/search, proteomes/search, uniprotkb/stream',
    'pdftotext (poppler): extract reference template paper text',
    'Python venv (numpy, scipy, biopython, openbabel, rdkit, pandas)',
    'OpenBabel (pybel / obabel CLI): 3D peptide conformation generation, PDB->PDBQT conversion',
    'RDKit: MolFromSequence rescue for Arg/His peptide NaN, MMFF optimization',
    'AutoDock Vina 1.2.5 (vina.exe): molecular docking',
    'awk / Python: PDB coordinate parsing, pocket-center extraction, receptor PDBQT header cleanup',
]:
    bullet(s)

# ===== VII. iDPPIV-SCM module =====
h1('VII. iDPPIV-SCM module - offline activity pre-screening (replaces PeptideRanker proxy scoring)')
para('Background and motivation', bold=True)
bullet('The PeptideRanker official server was long unavailable; and it is a "generic bioactivity" scorer, not optimal for DPP-4 specifically. The original proxy heuristic scored rich-Pro short peptides constantly at 1.000 with zero discrimination. Hence we switched to the DPP-IV-dedicated, fully-offline-reproducible iDPPIV-SCM (Scoring Card Method).')
para('Offline reproduction method', bold=True)
bullet('Dataset: downloaded the iDPPIV homologous benchmark from the public repo (WeiLab-BioChem/Structural-DPP-IV) - train 532+532, independent test 133+133, all standard 20 amino acids.')
bullet('Scoring card: locally recompute the global amino-acid-composition propensity score P(a)=log2(pos-frequency/neg-frequency); summing per-residue propensities over a peptide gives a continuous score (consistent with the paper abstract "propensity scores of amino acids"). Zero network, zero external dependency, fully reproducible.')
bullet('Interpretability: the learned propensities are biologically self-consistent - Pro scores +0.875 at the top (DPP-4 S1 pocket is specific to Pro; food-source DPP-4 inhibitory peptides are generally Pro-rich); Cys -2.482 strongly negative.')
para('Reproduction accuracy and key findings', bold=True)
t3 = doc.add_table(rows=1, cols=2)
t3.style = 'Light Grid Accent 1'
hh = t3.rows[0].cells
for i, t in enumerate(['Metric', 'Result']):
    hh[i].text = t
    hh[i].paragraphs[0].runs[0].bold = True
for r in [
    ('Independent test ACC', '~0.771 (literature reports ~0.797)'),
    ('Training-benchmark length confounding', 'positives mostly short peptides, negatives mostly long peptides/proteins; a pure "short=positive" baseline already reaches 0.820'),
    ('Implication for moso project', 'all candidates are 2-6 aa short peptides, length signal fails uniformly -> SCM score reflects genuine residue-composition signal'),
    ('iDPPIV score vs Vina dG correlation', 'Spearman rho = 0.067 (~0): activity propensity and binding free energy are orthogonal dimensions'),
]:
    c = t3.add_row().cells
    for i, v in enumerate(r):
        c[i].text = v

para('')
para('Funnel and docking comparison', bold=True)
bullet('Rescore all 4,950 2-6 aa short peptides (value range [-7.243, 4.431], mean -0.704, std 1.246, 68.7% predicted as DPP-IV inhibitory peptides), completely replacing the constant-1.000 proxy score.')
bullet('Feed the two-stage pipeline: 4,950 short peptides -> iDPPIV-prioritized candidates -> DPP4 structural-preference narrowing 565 -> Top-60 docking queue (overlaps 32 with the old proxy queue, adds 28 new).')
bullet('Fair re-docking (both old and new queues use the same RDKit MMFF94 pipeline, eliminating preparation confounding): at the set level the iDPPIV queue enriches stronger-binding peptides - dG<=-6.0 share 33.3% vs old 20.0%, and among the 32 overlapping peptides 21 are better in the new queue; but Delta=-0.35 kcal/mol falls within Vina\'s +/-0.5~1.0 noise band, so we do not overclaim "significantly better binding".')
bullet('Core conclusion: iDPPIV-SCM is an activity pre-screener (classification), not a binding-affinity predictor; together with Vina docking it forms a two-stage orthogonal filter. Its real value is upgrading the "down-server-dependent + zero-discrimination proxy" into a "literature-validated, offline-reproducible, interpretable" activity score - a hard methodological upgrade.')

para('Related deliverables', bold=True)
bullet('scripts/idppiv_scm/ (scoring-card model model.py + public dataset + validation/scoring scripts); data/moso_candidates_idppiv_short.tsv (4,950 short-peptide scores); docking/moso_dock_results_idppiv_clean.tsv, moso_dock_results_old_rdkit.tsv, moso_dock_compare.tsv; docs/methodology_replacement_report.md.')

out_path = os.path.join(ROOT, 'docs', 'moso_dpp4_peptide_project_summary.docx')
doc.save(out_path)
print('Saved:', out_path)
