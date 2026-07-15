# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- repo-relative paths (scripts/reports/ -> repo root) ---
ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

d = Document()
d.styles["Normal"].font.name = "Calibri"
d.styles["Normal"].font.size = Pt(10.5)

def h(t):
    p = d.add_heading(t, level=1); return p
def para(t, bold=False):
    p = d.add_paragraph(); r = p.add_run(t); r.bold = bold; return p

h("Moso Bamboo (Moso bamboo) Virtual Gastrointestinal Digestion Validation Report")
para("Species: Phyllostachys edulis (Moso bamboo) | UniProt taxonomy ID: 38705", True)
para("Purpose: verify that, as a novel DPP4 inhibitory-peptide study object, moso bamboo's UniProtKB protein pool can support the entire computational pipeline (vs the garlic template paper).")
para("Methodology basis: Cheng et al., Bioorganic Chemistry 175 (2026) 109801 (garlic DPP4 peptides, the methodological template for this work).")

h("1. Data source")
para("• Protein set: downloaded from UniProtKB /uniprotkb/stream all entries with taxonomy_id:38705, format=fasta.")
para("• Download result: 253 protein sequences, 131,713 bytes, all with real amino-acid sequences (no placeholders).")
para("• Sequence composition: average 408.7 aa/entry, total residues 103,409; representative proteins include cellulose synthase CesA, CDK, photosystem II CP43 and other real structural proteins.")
para("• Note: all 253 entries are TrEMBL (unreviewed) computational predictions, 0 entries Swiss-Prot (reviewed); the garlic template is also ~95% TrEMBL, acceptable in this field - disclose truthfully at submission.")

h("2. Virtual digestion rules (replicating the template paper's PeptideCutter settings)")
para("• pepsin (pH 1.3): cleavage at N-terminal side of F/Y/W/L")
para("• trypsin: cleavage at C-terminal side of K/R")
para("• chymotrypsin (specific): cleavage at C-terminal side of F/Y/W (L not cut; no cut before Pro)")
para("• Merge the three enzyme cut sites -> fragments; keep unique peptides of 2-20 aa and without X.")
para("(Note: this run replicates the above cut-site rules in code; cut counts match the garlic template paper; the formal manuscript must use the actual ExPASy PeptideCutter output as ground truth.)")

h("3. Digestion results")
t = d.add_table(rows=1, cols=3); t.style = "Light Grid Accent 1"
hdr = t.rows[0].cells
hdr[0].text = "Metric"; hdr[1].text = "Moso (this validation)"; hdr[2].text = "Garlic template paper"
rows = [
 ("Input protein count","253","113 (n=54 actually used for digestion)"),
 ("Total residues","103,409","-"),
 ("Unique peptides (all lengths, no X)","7,988","5,672"),
 ("Unique peptides 2-6 aa","4,950","1,442"),
 ("Unique peptides 2-20 aa","7,870","-"),
 ("PeptideRanker>0.5 candidates (est ~40%)","~1,980","249"),
 ("Final docking/synthesis count (pending)","-","34"),
]
for r0,r1,r2 in rows:
    c = t.add_row().cells; c[0].text = r0; c[1].text = r1; c[2].text = r2

h("4. Conclusion and pass/fail judgment")
para("OK: Passes, and outperforms the garlic template.", True)
para("• Moso 253 proteins yield 7,988 unique peptides, 4,950 2-6 aa short peptides - about 3.4x the scale of the garlic template (1,442 2-6 aa).")
para("• The protein pool (253 entries / 103k residues) is over 2x the garlic pool (113 entries), sufficient to support the 47-54 protein digestion subset and the full downstream funnel (PeptideRanker -> docking -> MD/MM-PBSA -> network pharmacology -> Caco-2).")
para("• The yam failed due to a data/species problem (D. polystachya only 20 entries -> only 237 peptides -> funnel collapse); moso has no such problem.")
para("WARNING: Honesty footnote: moso has 0 manually-reviewed entries, all TrEMBL predictions; must truthfully declare the data source and predictive nature in Methods.")
para("NEXT: Suggested next step: use PeptideRanker to score the 4,950 short peptides -> hydrophobicity/toxicity (AllerTOP/ToxinPred) filtering -> AutoDock Vina dock DPP4 (PDB 1WCY) -> in-vitro Gly-Pro-pNA activity validation, replicating the template's full pipeline.")

h("5. Deliverables")
para("• data/moso_253.fasta - moso 253 proteins FASTA (directly usable for PeptideCutter / re-run)")
para("• data/moso_253_peptides_strict.txt - 7,988 unique-peptide list")
para("• data/moso_253_peptides.txt - relaxed-rule 7,472 unique-peptide list (sensitivity control)")
para("• Script: scripts/rerun_digestion_moso253_strict.py")

_out = os.path.join(ROOT, "docs", "moso_digestion_validation_report.docx")
os.makedirs(os.path.dirname(_out), exist_ok=True)
d.save(_out)
print(f"saved {_out}")
