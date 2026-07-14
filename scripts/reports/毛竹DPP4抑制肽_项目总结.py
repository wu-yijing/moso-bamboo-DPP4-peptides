# -*- coding: utf-8 -*-
"""生成毛竹 DPP4 抑制肽项目总结 docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# 仓库根 = 本脚本(scripts/reports/)的上两级；使 clone 后即可独立生成
ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

doc = Document()

# 默认字号
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

def h1(t):
    p = doc.add_heading(t, level=1)
    return p

def h2(t):
    p = doc.add_heading(t, level=2)
    return p

def para(t, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = bold
    return p

def bullet(t):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(t)
    return p

def numbered(t):
    p = doc.add_paragraph(style='List Number')
    p.add_run(t)
    return p

# ===== 标题 =====
title = doc.add_heading('毛竹 DPP4 抑制肽项目：已完成工作总览', level=0)
sub = doc.add_paragraph('生成时间：2026-07-14 更新  |  仓库：github.com/wu-yijing/moso-bamboo-DPP4-peptides')
sub.runs[0].italic = True

# ===== 一、立项目标与转向原因 =====
h1('一、立项目标与转向原因')
para('原山药（Dioscorea polystachya）项目因 UniProt 仅 20 条 curated 蛋白、本机重跑虚拟消化只能产出 237 条肽（崩溃于原 5,230 漏斗），且 7 条实验验证肽（FWPQY 等）0/7 可由这 20 条蛋白切出 → 山药无法支撑 DPP4 声明。遂转向参考模板论文《Discovery of garlic-derived peptides as natural DPP4 inhibitors》（Cheng et al., Bioorganic Chemistry 175, 2026, 109801）方法，重选毛竹（Moso bamboo）为研究对象。')

# ===== 二、各步骤明细 =====
h1('二、各步骤明细（结果 / 软件 / 方法）')

# 步骤1
h2('步骤 1 — 物种鉴定与蛋白池核实')
bullet('结果：毛竹 = Phyllostachys edulis，UniProt taxonomy 38705；curated（Swiss-Prot）= 0 条，全部条目（含 TrEMBL）= 253 条，无参考蛋白质组；平均 409 aa，总残基 103,409。对比模板大蒜（A. sativum）仅 113 条池 → 毛竹池约为大蒜 2.2 倍，达标。')
bullet('软件：curl + UniProt REST API（taxonomy/search、uniprotkb/search、proteomes/search）。')
bullet('方法：先 taxonomy 反查物种编号（勿信俗名），再分别统计 reviewed vs all 条目数；基准对照模板论文实际用量。')

# 步骤2
h2('步骤 2 — 蛋白序列下载')
bullet('结果：moso_253.fasta（131 KB，253 条真实序列，0 占位符）。')
bullet('软件：curl + UniProtKB /stream（FASTA 格式）。')
bullet('方法：query=taxonomy_id:38705，全库条目池（非纯 reviewed，与模板论文实际做法一致）。')

# 步骤3
h2('步骤 3 — 模拟胃肠消化（虚拟酶解）')
bullet('结果：严格规则（chymotrypsin 切 F/Y/W）：唯一肽 7,988 条，其中 2–6 aa 短肽 4,950 条；宽松规则（加切 L）：唯一肽 7,472 条，2–6 aa 4,761 条。规模 = 大蒜模板（1,442 短肽）的 3.4 倍。')
bullet('软件：Python（venv：numpy/scipy/biopython）+ 自写脚本 rerun_digestion_moso253_strict.py。')
bullet('方法：复刻 ExPASy PeptideCutter —— pepsin (pH1.3) + trypsin + 特异性 chymotrypsin（切 F/Y/W 的 C 端，Pro 前不停），合并去重，保留 2–6 aa。')

# 步骤4
h2('步骤 4 — PeptideRanker 风格打分 + 过敏原/毒性过滤')
bullet('结果：>0.5 候选 4,333 条 → 去过敏原（AllerTOP 风格）/ 去毒（ToxinPred 风格）→ 4,289 条。')
bullet('软件：Python（自写代理启发式）。')
bullet('方法：早期采用透明代理打分（基于文献理化特征：N 端疏水、Pro 含量、分子量）。因 PeptideRanker 官方服务器长期不可用，且该代理分对富 Pro 短肽恒为 1.000、零区分力，已由「iDPPIV-SCM 本地离线复现」替代（见第七节），过敏/毒性层由 ToxinPred 3.0 + AlgPred 2.0 官方输出替代。')

# 步骤5
h2('步骤 5 — DPP4 结构偏好收窄 + 对接队列')
bullet('结果：3–5 aa、N 端疏水、第 2 位 Pro/Ala 优先 → 候选池 2,019 条；取优先 60 条进入对接队列。')
bullet('软件：Python 自写 moso_pipeline_filter2.py。')
bullet('方法：依据 DPP4 抑制肽已知结构规律（短肽、N 端疏水、Pro 富集）二次筛选。')

# 步骤6
h2('步骤 6 — 配体 3D 结构制备（PDBQT）')
bullet('结果：60/60 配体 PDBQT 全部生成成功（3–7 KB/个，零 NaN）。')
bullet('软件：OpenBabel（pybel / obabel CLI）+ RDKit（补救 8 条含 Arg/His 肽）。')
bullet('方法：主体 SMILES → make3D → PDBQT；8 条 Arg/His 肽因缺力场数据产生 NaN → 改用 RDKit MolFromSequence + AddHs + ETKDGv3 嵌入 + MMFFOptimizeMolecule 优化 → 转 PDBQT。')

# 步骤7
h2('步骤 7 — 受体与口袋准备')
bullet('结果：1WCY_receptor.pdbqt（12,248 原子，格式正确）；口袋盒子 moso_box.txt 中心 (62.8, 47.7, 4.8)，size 30³。')
bullet('软件：RCSB PDB 下载 + OpenBabel（obabel CLI）+ awk 固定列解析。')
bullet('方法：下载 DPP4 晶体 1WCY（西格列汀 sitagliptin 结合态）→ 转 PDBQT → 仅保留 ATOM/TER/END（剔除 HEADER 等头）→ 实测配体 A1201 坐标定网格中心（模板论文给的 54,62,37 为同口袋近似，以实测为准更稳）。')

# 步骤8
h2('步骤 8 — 分子对接（AutoDock Vina）')
para('结果：60/60 完成。Top 10（结合自由能 dG, kcal/mol）：', bold=True)

# 表格
tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Light Grid Accent 1'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
for i, t in enumerate(['排名', '肽序', 'dG (kcal/mol)', '类别']):
    hdr[i].text = t
    hdr[i].paragraphs[0].runs[0].bold = True
rows = [
    ('1', 'LPPQ (Leu-Pro-Pro-Gln)', '-7.472', '中-强'),
    ('2', 'APSPE (Ala-Pro-Ser-Pro-Glu)', '-7.150', '中-强'),
    ('3', 'LAPSP (Leu-Ala-Pro-Ser-Pro)', '-7.087', '中-强'),
    ('4', 'LPGP (Leu-Pro-Gly-Pro)', '-7.075', '中-强'),
    ('5', 'LPINP (Leu-Pro-Ile-Asn-Pro)', '-6.988', '中-强'),
    ('6', 'LPSP (Leu-Pro-Ser-Pro)', '-6.867', '中-强'),
    ('7', 'LPCPR (Leu-Pro-Cys-Pro-Arg)', '-6.835', '中-强'),
    ('8', 'LPGDP (Leu-Pro-Gly-Asp-Pro)', '-6.793', '中-强'),
    ('9', 'LPDDP (Leu-Pro-Asp-Asp-Pro)', '-6.693', '中-强'),
    ('10', 'APSQP (Ala-Pro-Ser-Gln-Pro)', '-6.515', '中-强'),
]
for r in rows:
    c = tbl.add_row().cells
    for i, v in enumerate(r):
        c[i].text = v

para('')
bullet('分布：中-强（-6.5 ~ -8）10 条 / 中等（-5 ~ -6.5）49 条 / 弱（>-5）1 条（CPPSK -4.856）；无强结合（<-8）。冠军 LPPQ（Leu-Pro-Pro-Gln，双 Pro）完全符合 DPP4 抑制肽文献特征，dG 量级与模板大蒜肽 WPHY/WPQY 相当。')
bullet('软件：AutoDock Vina 1.2.5（vina.exe，本机 Windows）。')
bullet('方法：--exhaustiveness 4 --cpu 2，单对接盒 30³，输出 9 个构象/肽，取最佳 mode 的 dG。')

# ===== 三、交付物清单 =====
h1('三、当前交付物清单（仓库内 data/ 与 docking/）')
for f in [
    'moso_253.fasta（253 蛋白）',
    'moso_253_peptides_strict.txt / _peptides.txt（唯一肽清单）',
    'moso_candidates_pr_filtered.txt（2,019 候选）',
    'moso_dock_queue.txt（60 对接队列）',
    '1WCY_receptor.pdbqt、moso_box.txt、vina.exe',
    'moso_ligands/（60 配体 PDBQT + 60 对接输出构象）',
    'moso_dock_results.tsv、moso_dock_ranking.txt（完整排名）',
]:
    bullet(f)

# ===== 四、待完成 =====
h1('四、尚待完成的后续步骤（按模板全程）')
t2 = doc.add_table(rows=1, cols=3)
t2.style = 'Light Grid Accent 1'
h = t2.rows[0].cells
for i, t in enumerate(['阶段', '状态', '说明']):
    h[i].text = t
    h[i].paragraphs[0].runs[0].bold = True
for r in [
    ('官方 PeptideRanker/AllerTOP/ToxinPred 验证', '未做', '当前为代理启发式，须替换'),
    ('MD + MM/PBSA（GROMACS）', '未做', '模板用 50→150 ns 验证 Top 肽'),
    ('网络药理学', '未做', 'SwissTargetPrediction→STRING→DAVID'),
    ('体外活性（Gly-Pro-pNA + Caco-2）', '未做', '真实活性终点，须实验'),
]:
    c = t2.add_row().cells
    for i, v in enumerate(r):
        c[i].text = v

# ===== 五、局限 =====
h1('五、必须如实披露的局限')
bullet('活性预筛打分已由 iDPPIV-SCM 本地离线复现替代原代理启发式（文献验证、可离线复现、可解释）；过敏/毒性层用 ToxinPred 3.0 + AlgPred 2.0 官方输出。iDPPIV-SCM 训练基准集存在长度混杂（负样本多为长肽/蛋白），故仅用作候选排序/软过滤，不作确定性活性判定。')
bullet('对接为静态 dG 估计，需 MD/MM-PBSA + 体外实验确认，不能直接当活性结论。')
bullet('毛竹 0 条人工审阅条目（全 TrEMBL 预测），Methods 须如实披露来源与预测性质（大蒜模板同样 ~95% 为 TrEMBL，可接受）。')
bullet('WPHY/WPQY/VAPGW 是大蒜肽，不应用于毛竹真值校验。')

# ===== 六、软件栈 =====
h1('六、本项目软件栈汇总')
for s in [
    'UniProt REST API（curl）：taxonomy/search、uniprotkb/search、proteomes/search、uniprotkb/stream',
    'pdftotext（poppler）：抽取参考模板论文文本',
    'Python venv（numpy、scipy、biopython、openbabel、rdkit、pandas）',
    'OpenBabel（pybel / obabel CLI）：3D 肽构象生成、PDB→PDBQT 转换',
    'RDKit：MolFromSequence 补救 Arg/His 肽 NaN、MMFF 优化',
    'AutoDock Vina 1.2.5（vina.exe）：分子对接',
    'awk / Python：PDB 坐标解析、口袋中心提取、受体 PDBQT 头文件清理',
]:
    bullet(s)

# ===== 七、iDPPIV-SCM 模块 =====
h1('七、iDPPIV-SCM 模块 — 离线活性预筛（替代 PeptideRanker 代理打分）')
para('背景与动机', bold=True)
bullet('PeptideRanker 官方服务器长期不可用；且其为「通用生物活性」打分器，对 DPP-4 专项并非最优。原代理启发式分对富 Pro 短肽恒为 1.000、零区分力。故改用 DPP-IV 专用、可完全离线复现的 iDPPIV-SCM（Scoring Card Method）。')
para('离线复现方法', bold=True)
bullet('数据集：从公开同源仓库（WeiLab-BioChem/Structural-DPP-IV）下载 iDPPIV 同源基准集——train 532+532、独立测试 133+133，均为标准 20 种氨基酸。')
bullet('评分卡：本地重算全局氨基酸组成型倾向性得分 P(a)=log2(正样本频率/负样本频率)，对一条肽求 Σ 残基倾向性得到连续分（与论文摘要「propensity scores of amino acids」一致）。零联网、零外部依赖、完全可复现。')
bullet('可解释性：学到的倾向性生物学自洽——Pro 得分 +0.875 居首（DPP-4 的 S1 口袋对 Pro 特异，食源 DPP-4 抑制肽普遍富 Pro）；Cys −2.482 强烈负向。')

para('复现精度与关键发现', bold=True)
t3 = doc.add_table(rows=1, cols=2)
t3.style = 'Light Grid Accent 1'
hh = t3.rows[0].cells
for i, t in enumerate(['指标', '结果']):
    hh[i].text = t
    hh[i].paragraphs[0].runs[0].bold = True
for r in [
    ('独立测试 ACC', '≈0.771（文献报告 ≈0.797）'),
    ('训练基准集长度混杂', '正样本多为短肽、负样本多为长肽/蛋白；纯「短=阳性」基线即达 0.820'),
    ('对毛竹项目的意义', '所有候选均为 2–6 aa 短肽，长度信号一致失效 → SCM 分反映真实残基组成信号'),
    ('iDPPIV 分 vs Vina dG 相关性', 'Spearman ρ = 0.067（≈0）：活性倾向与结合自由能是正交维度'),
]:
    c = t3.add_row().cells
    for i, v in enumerate(r):
        c[i].text = v

para('')
para('筛选漏斗与对接对比', bold=True)
bullet('对全部 4,950 条 2–6 aa 短肽重新打分（取值域 [−7.243, 4.431]，均值 −0.704、标准差 1.246，68.7% 预测为 DPP-IV 抑制肽），彻底替换恒为 1.000 的代理分。')
bullet('接入两级流程：4,950 短肽 → iDPPIV 优先化候选 → DPP4 结构偏好收窄 565 → Top-60 对接队列（与旧代理队列重叠 32、新增 28）。')
bullet('公平重对接（新旧队列均用同一 RDKit MMFF94 管线，消除制备混杂）：集合层面 iDPPIV 队列富集更强结合肽——dG≤−6.0 占比 33.3% vs 旧 20.0%，重叠 32 肽中新队列 21 个更优；但 Δ=−0.35 kcal/mol 落在 Vina ±0.5~1.0 噪声带内，不夸大为「结合显著更优」。')
bullet('核心结论：iDPPIV-SCM 是活性预筛器（分类），非结合亲和力预测器；与 Vina 对接构成两级正交过滤。其真正价值是把「宕机依赖 + 零区分力代理」升级为「文献验证、可离线复现、可解释」的活性评分——方法学硬升级。')

para('相关产物', bold=True)
bullet('scripts/idppiv_scm/（评分卡模型 model.py + 公开数据集 + 验证/打分脚本）；data/moso_candidates_idppiv_short.tsv（4,950 短肽评分）；docking/moso_dock_results_idppiv_clean.tsv、moso_dock_results_old_rdkit.tsv、moso_dock_compare.tsv；docs/方法学替换报告.md。')

out_path = os.path.join(ROOT, 'docs', '毛竹DPP4抑制肽_项目总结.docx')
doc.save(out_path)
print('Saved:', out_path)
