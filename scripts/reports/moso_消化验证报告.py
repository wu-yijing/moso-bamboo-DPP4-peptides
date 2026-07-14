# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- repo-relative paths (scripts/reports/ -> repo root) ---
ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

d = Document()
d.styles["Normal"].font.name = "Calibri"
d.styles["Normal"].font.size = Pt(10.5)

def h(t): 
    p=d.add_heading(t,level=1); return p
def para(t,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=bold; return p

h("毛竹 (Moso bamboo) 虚拟胃肠消化验证报告")
para("物种: Phyllostachys edulis (Moso bamboo) | UniProt taxonomy ID: 38705", True)
para("生成目的: 核实毛竹作为 DPP4 抑制肽新研究对象，其 UniProtKB 蛋白池能否支撑整条计算管线（对照大蒜模板论文）。")
para("方法学依据: Cheng et al., Bioorganic Chemistry 175 (2026) 109801（大蒜 DPP4 肽，本工作方法学模板）。")

h("1. 数据来源")
para("• 蛋白集: 通过 UniProtKB /uniprotkb/stream 下载 taxonomy_id:38705 全部条目，format=fasta。")
para("• 下载结果: 253 条蛋白序列，131,713 字节，全部含真实氨基酸序列（无占位符）。")
para("• 序列构成: 平均 408.7 aa/条，总残基 103,409；代表蛋白包括纤维素合酶 CesA、CDK、光系统 II CP43 等真实结构蛋白。")
para("• 注: 253 条均为 TrEMBL(unreviewed)计算预测，0 条 Swiss-Prot(reviewed)；大蒜模板同样 ~95% 为 TrEMBL，本领域可接受，投稿时如实披露即可。")

h("2. 虚拟消化规则（复刻模板论文 PeptideCutter 设定）")
para("• pepsin (pH 1.3): 酶切位点 F/Y/W/L 的 N 端侧")
para("• trypsin: 酶切位点 K/R 的 C 端侧")
para("• chymotrypsin (specific): 酶切位点 F/Y/W 的 C 端侧（L 不切；Pro 前不停）")
para("• 合并三酶切位点 -> 片段；保留 2–20 aa 且不含 X 的唯一肽。")
para("（说明: 本运行以代码复刻上述切位规则，切点数与大葱模板论文所述一致；正式稿件须以 ExPASy PeptideCutter 实际输出为准。）")

h("3. 消化结果")
t = d.add_table(rows=1, cols=3); t.style="Light Grid Accent 1"
hdr=t.rows[0].cells
hdr[0].text="指标"; hdr[1].text="毛竹 (本验证)"; hdr[2].text="大蒜模板论文"
rows=[
 ("输入蛋白条数","253","113 (n=54 实际用于消化)"),
 ("总残基","103,409","—"),
 ("唯一肽（全部长度, 无X）","7,988","5,672"),
 ("唯一肽 2–6 aa","4,950","1,442"),
 ("唯一肽 2–20 aa","7,870","—"),
 ("PeptideRanker>0.5 候选（估 ~40%）","~1,980","249"),
 ("最终对接/合成数（待跑）","—","34"),
]
for r0,r1,r2 in rows:
    c=t.add_row().cells; c[0].text=r0; c[1].text=r1; c[2].text=r2

h("4. 结论与达标判定")
para("✅ 达标，且优于大蒜模板。", True)
para("• 毛竹 253 条蛋白产出 7,988 条唯一肽、4,950 条 2–6 aa 短肽，规模是大蒜模板（1,442 条 2–6 aa）的 ~3.4 倍。")
para("• 蛋白池（253 条 / 103k 残基）是大蒜池（113 条）的 2 倍以上，足以支撑 47–54 蛋白的消化子集与完整下游漏斗（PeptideRanker → 对接 → MD/MM-PBSA → 网络药理学 → Caco-2）。")
para("• 山药失败是数据/物种问题（D. polystachya 仅 20 条 → 仅 237 肽 → 漏斗崩溃）；毛竹不存在此问题。")
para("⚠️ 诚实注脚: 毛竹 0 条人工审阅条目，全部为 TrEMBL 预测；须在 Methods 中如实声明数据来源与预测性质。")
para("▶ 下一步建议: 用 PeptideRanker 对 4,950 条短肽打分 → 疏水性/毒性(AllerTOP/ToxinPred)过滤 → AutoDock Vina 对接 DPP4 (PDB 1WCY) → 体外 Gly-Pro-pNA 活性验证，复刻模板全流程。")

h("5. 交付文件")
para("• data/moso_253.fasta — 毛竹 253 条蛋白 FASTA（可直接用于 PeptideCutter / 重跑）")
para("• data/moso_253_peptides_strict.txt — 7,988 条唯一肽清单")
para("• data/moso_253_peptides.txt — 宽松规则 7,472 条唯一肽清单（敏感性对照）")
para("• 脚本: scripts/rerun_digestion_moso253_strict.py")

_out = os.path.join(ROOT, "docs", "毛竹_消化验证报告.docx")
os.makedirs(os.path.dirname(_out), exist_ok=True)
d.save(_out)
print(f"saved {_out}")
